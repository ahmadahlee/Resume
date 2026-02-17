from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

def get_json_data(filename):
    with open(filename, 'r') as f:
        return json.load(f)

def create_resume(json_file):
    data = get_json_data(json_file)
    basics = data.get('basics', {})
    
    document = Document()

    # Style configuration
    style = document.styles['Normal']
    font = style.font
    # Trying to match 'EB Garamond' with common serif fonts
    font.name = 'Garamond' 
    font.size = Pt(11)

    # Name
    name_paragraph = document.add_paragraph()
    name_run = name_paragraph.add_run(basics.get('name', ''))
    name_run.bold = False # Theme is not bolded in HTML
    name_run.font.size = Pt(30)
    name_run.font.color.rgb = RGBColor(0, 0, 0) 
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_paragraph.paragraph_format.space_before = Pt(20)

    # Title
    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run(basics.get('label', ''))
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(85, 85, 85) # Medium gray
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(10)

    # Contact Info
    contact_paragraph = document.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info = []
    if basics.get('phone'): contact_info.append(basics['phone'])
    if basics.get('email'): contact_info.append(basics['email'])
    for p in basics.get('profiles', []):
        contact_info.append(p.get('url', '').replace('https://', ''))
    
    run = contact_paragraph.add_run(' | '.join(contact_info))
    run.font.size = Pt(10)
    contact_paragraph.paragraph_format.space_after = Pt(20)

    # Summary
    if basics.get('summary'):
        p = document.add_paragraph(basics['summary'])
        p.paragraph_format.space_after = Pt(12)

    # Helper for Section Headers
    def add_section_header(text):
        p = document.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        # Horizontal line effectively simulated by header styling
        p.paragraph_format.keep_with_next = True

    # Experience
    if data.get('work'):
        add_section_header("Experience")
        for job in data['work']:
            p = document.add_paragraph()
            run = p.add_run(f"{job.get('company', job.get('name', ''))}")
            run.bold = True
            
            # Tab for date alignment
            tab_stop = Pt(450)
            p.paragraph_format.tab_stops.add_tab_stop(tab_stop, alignment=2)
            
            end_date = job.get('endDate')
            if not end_date:
                end_date = "Present"
            
            p.add_run(f"\t{job.get('startDate', '')} - {end_date}").font.italic = True
            
            p2 = document.add_paragraph()
            p2.add_run(job.get('position', '')).font.italic = True
            p2.paragraph_format.space_after = Pt(2)

            if job.get('summary'):
                 document.add_paragraph(job['summary'])

            if job.get('highlights'):
                for h in job['highlights']:
                    p = document.add_paragraph(h, style='List Bullet')
                    p.paragraph_format.left_indent = Pt(20)
            
            # Spacer between different jobs
            document.add_paragraph().paragraph_format.space_after = Pt(8)
        
        # Larger spacer at the end of Experience section
        document.add_paragraph().paragraph_format.space_after = Pt(15)

    # Projects (Shared with Work in some themes, but checking explicitly)
    if data.get('projects'):
        add_section_header("Projects")
        for project in data['projects']:
            p = document.add_paragraph()
            p.add_run(project.get('name', '')).bold = True
            
            p.paragraph_format.tab_stops.add_tab_stop(Pt(450), alignment=2)
            p.add_run(f"\t{project.get('startDate', '')} - {project.get('endDate', 'Present')}").font.italic = True
            
            if project.get('description'):
                document.add_paragraph(project['description'])
            
            if project.get('highlights'):
                for h in project['highlights']:
                    p = document.add_paragraph(h, style='List Bullet')
                    p.paragraph_format.left_indent = Pt(20)
            
            document.add_paragraph().paragraph_format.space_after = Pt(8)
        
        document.add_paragraph().paragraph_format.space_after = Pt(15)

    # Education Section
    if data.get('education'):
        add_section_header("Education")
        for edu in data['education']:
            p = document.add_paragraph()
            p.add_run(edu.get('institution', '')).bold = True
            
            p.paragraph_format.tab_stops.add_tab_stop(Pt(450), alignment=2)
            p.add_run(f"\t{edu.get('startDate', '')} - {edu.get('endDate', '')}").font.italic = True
            
            p2 = document.add_paragraph(edu.get('studyType', '') + " " + edu.get('area', ''))
            p2.paragraph_format.space_after = Pt(12)
        
        document.add_paragraph().paragraph_format.space_after = Pt(15)

    # Awards Section
    if data.get('awards'):
        add_section_header("Awards")
        for award in data['awards']:
            p = document.add_paragraph()
            p.add_run(award.get('title', '')).bold = True
            p.add_run(f" - {award.get('awarder', '')}")
            
            if award.get('summary'):
                document.add_paragraph(award['summary'])
            document.add_paragraph().paragraph_format.space_after = Pt(8)
        
        document.add_paragraph().paragraph_format.space_after = Pt(15)

    # Skills
    if data.get('skills'):
        add_section_header("Skills")
        for skill in data['skills']:
            p = document.add_paragraph()
            p.add_run(f"{skill.get('name')}: ").bold = True
            p.add_run(", ".join(skill.get('keywords', [])))
            p.paragraph_format.space_after = Pt(4)
        
        document.add_paragraph().paragraph_format.space_after = Pt(15)

    # Interests
    if data.get('interests'):
        add_section_header("Interests")
        for item in data['interests']:
            p = document.add_paragraph()
            p.add_run(f"{item.get('name')}: ").bold = True
            p.add_run(", ".join(item.get('keywords', [])))
            p.paragraph_format.space_after = Pt(4)
        
        document.add_paragraph().paragraph_format.space_after = Pt(15)

    # Languages
    if data.get('languages'):
        add_section_header("Languages")
        for lang in data['languages']:
            p = document.add_paragraph()
            p.add_run(f"{lang.get('language')}: ").bold = True
            p.add_run(lang.get('fluency', ''))
            p.paragraph_format.space_after = Pt(4)

    # Save
    document.save('ahmad_ali_resume.docx')
    print("Resume saved as ahmad_ali_resume.docx")

if __name__ == "__main__":
    # If run via automate.py from root, path remains the same
    create_resume('resume.json')
